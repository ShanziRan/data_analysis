from pathlib import Path
import math

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("output/hitl/OCR_Human_Review_Model_Leader_Brief.docx")
ASSET_DIR = Path("output/hitl/doc_assets")


def diagram_font(size, bold=False):
    choices = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    for path in choices:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def multiline_center(draw, box, text, size=25, bold=False):
    font = diagram_font(size, bold)
    left, top, right, bottom = box
    lines = text.split("\n")
    heights = []
    for line in lines:
        bounds = draw.textbbox((0, 0), line, font=font)
        heights.append(bounds[3] - bounds[1])
    total = sum(heights) + (len(lines) - 1) * 6
    y = top + (bottom - top - total) / 2
    for line, height in zip(lines, heights):
        bounds = draw.textbbox((0, 0), line, font=font)
        width = bounds[2] - bounds[0]
        draw.text((left + (right - left - width) / 2, y), line, font=font, fill="black")
        y += height + 6


def arrow(draw, start, end):
    draw.line((start, end), fill="black", width=4)
    x, y = end
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 17
    left = (x - size * math.cos(angle - .5), y - size * math.sin(angle - .5))
    right = (x - size * math.cos(angle + .5), y - size * math.sin(angle + .5))
    draw.polygon([(x, y), left, right], fill="black")


def diagram_box(draw, box, label, fill="white", size=23, bold=False):
    draw.rounded_rectangle(box, radius=12, outline="black", width=3, fill=fill)
    multiline_center(draw, box, label, size, bold)


def training_diagram(path):
    image = Image.new("RGB", (2100, 380), "white")
    draw = ImageDraw.Draw(image)
    boxes = [(20 + i * 295, 90, 250 + i * 295, 275) for i in range(7)]
    labels = [
        "Reviewed data\nhuman.csv",
        "Label OCR change\nCaptured != Published",
        "Learn character\nconfusions",
        "Calculate 16\nOCR features",
        "Split by scan_id\ntrain / validation",
        "Train OCR and\ncorrection-type models",
        "Choose recall\nthreshold; refit\nand save",
    ]
    for box, label in zip(boxes, labels):
        diagram_box(draw, box, label, size=22)
    for left, right in zip(boxes, boxes[1:]):
        arrow(draw, (left[2] + 7, 182), (right[0] - 7, 182))
    image.save(path)


def scoring_diagram(path):
    image = Image.new("RGB", (2100, 980), "white")
    draw = ImageDraw.Draw(image)
    source = (700, 30, 1400, 150)
    prep = (700, 205, 1400, 335)
    diagram_box(draw, source, "Unreviewed OCR data: no_human.csv", size=25)
    diagram_box(draw, prep, "Parse answer key independently; calculate shared evidence", size=24)
    arrow(draw, (1050, 150), (1050, 205))

    ocr_head = (80, 410, 560, 530)
    ocr_model = (80, 590, 560, 735)
    ocr_out = (80, 795, 560, 925)
    ak_head = (650, 410, 1130, 530)
    ak_sem = (650, 590, 1130, 735)
    ak_gate = (1230, 590, 2020, 735)
    ak_out = (1230, 795, 2020, 925)
    combined = (650, 795, 1130, 925)

    diagram_box(draw, ocr_head, "Decision 1: OCR review", fill="#EAF2F8", size=25, bold=True)
    diagram_box(draw, ocr_model, "16 features -> logistic probability\n-> saved threshold and risk", size=23)
    diagram_box(draw, ocr_out, "requires_ocr_review\ncorrection type", fill="#F2F8FC", size=23)
    diagram_box(draw, ak_head, "Decision 2: AK coverage", fill="#EAF2F8", size=25, bold=True)
    diagram_box(draw, ak_sem, "Eligible non-MCQ answer\nwhole-string semantic vs each AK variation", size=22)
    diagram_box(draw, ak_gate, "Require strong semantic score + low surface match + gap >= 0.15\nBlock single-token, number, date/time, negation, polarity, unit and key-term conflicts", size=20)
    diagram_box(draw, ak_out, "possible_gap_suggestion\nconflict reasons and best variation", fill="#F2F8FC", size=22)
    diagram_box(draw, combined, "Combined routing\nrequires_any_human_review + review_reasons", fill="#F7F7F7", size=22)

    arrow(draw, (900, 335), (320, 410))
    arrow(draw, (1200, 335), (890, 410))
    arrow(draw, (320, 530), (320, 590))
    arrow(draw, (320, 735), (320, 795))
    arrow(draw, (890, 530), (890, 590))
    arrow(draw, (1130, 662), (1230, 662))
    arrow(draw, (1625, 735), (1625, 795))
    arrow(draw, (560, 860), (650, 860))
    arrow(draw, (1230, 860), (1130, 860))

    footer_font = diagram_font(22)
    footer = "Write flagged CSV and scoring_report.json with parameters, counts, rates and score distributions"
    bounds = draw.textbbox((0, 0), footer, font=footer_font)
    draw.text(((2100 - (bounds[2] - bounds[0])) / 2, 946), footer, font=footer_font, fill="black")
    image.save(path)


def set_alt_text(shape, description):
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def font(run, size=10.5, bold=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def para(doc, text, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    font(p.add_run(text))
    return p


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), size=12, bold=True)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.32)
    p.paragraph_format.first_line_indent = Inches(-.18)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.04
    font(p.add_run(text))
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(.32)
    p.paragraph_format.first_line_indent = Inches(-.18)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.04
    font(p.add_run(text))
    return p


OUT.parent.mkdir(parents=True, exist_ok=True)
training_path = OUT.parent / "training_workflow_current.png"
scoring_path = OUT.parent / "scoring_workflow_current.png"
training_diagram(training_path)
scoring_diagram(scoring_path)

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = Inches(.7)
section.left_margin = section.right_margin = Inches(.8)
section.header_distance = section.footer_distance = Inches(.3)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor(0, 0, 0)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.08

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(2)
font(title.add_run("OCR review and answer-key coverage pipeline"), size=18, bold=True)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(8)
font(subtitle.add_run("Concise discussion guide: current training, scoring and review workflow"), size=11)

heading(doc, "Purpose")
para(doc, "Prioritise OCR answers for human review while separately identifying possible valid answers missing from the answer key. Both outputs are suggestions for reviewers; neither changes an answer automatically.")

heading(doc, "Two independent decisions")
bullet(doc, "OCR review: predicts whether Captured is likely to need correction, using patterns learned from human-reviewed data.")
bullet(doc, "Answer-key coverage: uses whole-string semantic comparison and conservative conflict rules to suggest possible missing AK variations.")
bullet(doc, "Combined routing: requires_any_human_review joins the queues while review_reasons preserves why each row was flagged.")

heading(doc, "OCR training data")
bullet(doc, "Source: data/hitl/human.csv - 341,190 reviewed rows.")
bullet(doc, "Captured is the original OCR answer; Published is the human-verified answer.")
bullet(doc, "OCR label: review needed when normalised Captured differs from Published; otherwise no correction.")
bullet(doc, "7,857 rows were corrected (2.3%); class weights compensate for the imbalance.")
bullet(doc, "The AK-gap rule is not trained because accepted/rejected AK-suggestion feedback is not yet available.")

heading(doc, "Training workflow")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(3)
shape = p.add_run().add_picture(str(training_path), width=Inches(6.75))
set_alt_text(shape, "Current OCR training workflow from reviewed data to saved models and recall threshold")

heading(doc, "Latest validation snapshot")
bullet(doc, "Split by scan_id so related rows remain in one fold.")
bullet(doc, "Target recall 98%; measured validation recall 98.03% at threshold 0.6814.")
bullet(doc, "Validation precision 24.95%, review rate 8.84%, and average precision 0.6127.")

heading(doc, "Data rules")
bullet(doc, "Published creates labels but is never a scoring feature.")
bullet(doc, "AK parsing is independent of Captured and bounded to 100 variations by default.")
bullet(doc, "Blank non-MCQ captures have no distance or similarity; is_blank carries the state.")

doc.add_page_break()

heading(doc, "Evidence used by the OCR model")
for item in [
    "ocr_confidence - OCR engine confidence.",
    "edit_similarity - normalised similarity to the closest accepted variation.",
    "char_ngram_similarity - overlap between short character sequences.",
    "token_similarity - exact word-set overlap; it is not semantic embedding.",
    "answer_exact, is_blank and is_mcq - structural indicators.",
    "captured_length, best_answer_length and length_difference.",
    "digit_fraction, alpha_fraction and punctuation_fraction.",
    "variation_count - bounded accepted variations considered.",
    "known_confusion_score and known_confusion_fraction - evidence from historical character changes.",
]:
    bullet(doc, item)
para(doc, "min_distance is logged as evidence but is not a model feature. It remains empty for blank non-MCQ captures.", after=4)

heading(doc, "Scoring workflow")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
shape = p.add_run().add_picture(str(scoring_path), width=Inches(6.75))
set_alt_text(shape, "Current scoring workflow with separate OCR review and answer-key coverage decisions and combined routing")

heading(doc, "OCR outputs")
bullet(doc, "ocr_review_probability and requires_ocr_review come from the trained model and saved threshold.")
bullet(doc, "ocr_risk_label: low below threshold; medium from threshold to max(0.8, threshold); high above that boundary.")
bullet(doc, "predicted_ocr_correction_type: no_correction, substitution, insertion, deletion, transposition, multiple_edits, blank_to_text, text_to_blank, case_or_whitespace or mcq_correction.")

doc.add_page_break()

heading(doc, "Answer-key coverage decision")
bullet(doc, "Eligible rows are non-MCQ, non-blank and not an exact accepted-answer match.")
bullet(doc, "all-MiniLM-L6-v2 embeds the complete Captured string and every complete AK variation; the highest cosine similarity is retained.")
bullet(doc, "possible_gap_suggestion requires semantic similarity at least 0.78 for single-word-versus-phrase or 0.80 for multi-word pairs, surface similarity at most 0.60, and semantic-surface gap at least 0.15.")
bullet(doc, "Conflicts block different single words, numbers, dates/times, negation, polarity, units or a changed key term in a short near-duplicate.")
bullet(doc, "Blocked rows keep semantic_similarity, best_semantic_variation and ak_conflict_reasons for audit.")

heading(doc, "AK labels")
bullet(doc, "possible_answer_key_gap - passes all semantic, surface-gap and conflict gates.")
bullet(doc, "blocked_by_conflict - semantically related but materially inconsistent.")
bullet(doc, "semantic_below_threshold, surface_match_existing_ak or insufficient_semantic_surface_gap - fails an evidence gate.")
bullet(doc, "semantic_not_scored - AK coverage was intentionally skipped.")

heading(doc, "Using the review queue")
bullet(doc, "Route requires_any_human_review = True; use review_reasons to separate OCR correction, possible AK gap or both.")
bullet(doc, "For OCR review, prioritise high then medium risk and show the scan plus similarity evidence.")
bullet(doc, "For AK review, show Captured, best semantic variation, semantic and surface scores, and conflict reasons.")
bullet(doc, "Record reviewer outcomes in new fields without overwriting Captured; keep a low-risk quality-control sample.")

heading(doc, "Run audit and latest scoring snapshot")
bullet(doc, "Each run writes the flagged CSV and scoring_report.json with model hash, parameters, thresholds, counts, rates and distributions.")
bullet(doc, "Latest run: 138,810 rows; 5,477 OCR flags, 193 AK-gap flags, and 5,501 combined review rows (3.96%).")
bullet(doc, "Recall and precision require verified Published truth and explicit --evaluate-with-published.")

heading(doc, "Points to agree with the lead")
bullet(doc, "Acceptable missed-correction rate and available reviewer capacity.")
bullet(doc, "Quality-control sample for low-risk and conflict-blocked AK rows.")
bullet(doc, "Reviewer labels to collect so AK-gap logic can later be trained and validated.")
bullet(doc, "Ownership of threshold changes, model versions and answer-key updates.")

doc.save(OUT)
print(OUT.resolve())
print(training_path.resolve())
print(scoring_path.resolve())
