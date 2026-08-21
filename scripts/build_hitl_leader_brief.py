from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("output/hitl/OCR_Human_Review_Model_Leader_Brief.docx")
BLUE = "2E74B5"
DARK = "1F4D78"
INK = "263238"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "E8F1F8"
PALE_GOLD = "FFF4D6"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)


def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, bold_lead=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True)
        set_font(p.add_run(text[len(bold_lead):]))
    else:
        set_font(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.5)
    p.paragraph_format.first_line_indent = Inches(-.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.10
    set_font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(.5)
    p.paragraph_format.first_line_indent = Inches(-.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.10
    set_font(p.add_run(text))
    return p


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_callout(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    set_font(p.add_run(label + "  "), bold=True, color=DARK)
    set_font(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for i, header in enumerate(headers):
        shade(table.rows[0].cells[i], LIGHT)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_font(p.add_run(header), bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_font(p.add_run(str(value)), size=10.2)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_break(doc):
    doc.add_page_break()


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(header.add_run("OCR HUMAN-REVIEW MODEL  |  LEADERSHIP BRIEF"), size=9, bold=True, color=MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(footer.add_run("Internal working document  •  20 August 2026"), size=9, color=MUTED)

# Opening masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(4)
set_font(p.add_run("LEADERSHIP BRIEF"), size=11, bold=True, color=BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
set_font(p.add_run("OCR Human-Review Prioritisation"), size=25, bold=True, color=DARK)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
set_font(p.add_run("How reviewed answers train the model, how new answers are scored, and how people remain in control"), size=13, color=MUTED)

for label, value in (
    ("Purpose", "Reduce unnecessary manual checking while directing reviewers to the rows most likely to contain OCR errors."),
    ("Current data", "341,190 human-reviewed rows; 7,857 changed by a reviewer (approximately 2.3%)."),
    ("Decision supported", "Which unreviewed OCR rows should enter a human-review queue first."),
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(label + ": "), bold=True)
    set_font(p.add_run(value))

heading(doc, "Executive summary", 1)
add_callout(doc, "Key message", "The model does not correct answers or replace reviewers. It estimates the likelihood that a reviewer would change an OCR-captured answer, then helps prioritise human attention.")
add_bullet(doc, "Training uses past human decisions: Captured is the OCR value; Published is the human-verified outcome.")
add_bullet(doc, "Scoring uses only information available before review. Published is never used as an input when predicting new rows.")
add_bullet(doc, "Each row receives a review probability, a yes/no review flag, a risk band, supporting similarity evidence, and a likely correction type.")
add_bullet(doc, "The review threshold is chosen to target high recall (95% by default), accepting a larger review queue in exchange for fewer missed corrections.")

heading(doc, "What success looks like", 1)
add_text(doc, "The central operational measure is: how many genuine OCR corrections are found for the amount of human review required. Ordinary accuracy is misleading because approximately 97.7% of reviewed rows were unchanged.")

add_page_break(doc)
heading(doc, "1. What the model is trained on", 1)
add_text(doc, "The supervised training source is data/hitl/human.csv. Every row has an original OCR result and a human-verified result.")
add_table(doc, ["Field", "Role in training", "Used when scoring?"], [
    ("Captured", "Original OCR-captured response", "Yes"),
    ("Published", "Human-verified response; creates the training label", "No"),
    ("ANSWER Key", "Defines acceptable answer variations", "Yes"),
    ("confidence", "OCR engine confidence", "Yes"),
    ("scan_id", "Keeps complete scans together during validation", "Operational only"),
], [1800, 4660, 2900])

heading(doc, "How the training label is created", 2)
add_text(doc, "After trimming surrounding spaces and ignoring letter case:")
add_table(doc, ["Observed outcome", "Training label"], [
    ("Captured equals Published", "No review needed / no correction"),
    ("Captured differs from Published", "Review needed / human correction"),
], [4680, 4680])
add_callout(doc, "Why weighting is needed", "Only about 2.3% of rows were corrected. The model gives corrected rows more influence during training so it cannot achieve a misleadingly high score by predicting ‘no review’ for everything.", PALE_GOLD)

heading(doc, "How validation stays honest", 2)
add_text(doc, "The data is split by scan_id: all answers from one scan remain either in training or in validation. This reduces leakage from related answers and gives a more realistic estimate of performance on future scans.")
add_text(doc, "During validation, common character confusions are learned only from the training portion. Once evaluation is complete, the deployment model is refitted using all reviewed rows.")

add_page_break(doc)
heading(doc, "2. Signals used by the model", 1)
doc.add_paragraph().paragraph_format.space_after = Pt(0)
add_text(doc, "The model combines several weak signals. No single similarity threshold decides the outcome.")
add_table(doc, ["Signal", "Plain-language meaning"], [
    ("OCR confidence", "How confident the OCR engine was in the captured text."),
    ("Edit similarity", "How few character insertions, deletions or substitutions separate Captured from the nearest accepted variation."),
    ("Character n-gram similarity", "How similar short overlapping character sequences are; useful for misspellings, names and short text."),
    ("Token similarity", "How much the set of words overlaps with an accepted variation."),
    ("Exact answer match", "Whether Captured exactly matches an accepted answer variation."),
    ("Text shape", "Blank/MCQ status, length, digits, letters and punctuation."),
    ("Known confusions", "Whether required edits resemble character changes repeatedly made by reviewers in the past."),
], [2700, 6660])

heading(doc, "Special handling", 2)
add_bullet(doc, "Blank non-MCQ capture: no distance or similarity is calculated. The dedicated blank indicator is used instead.")
add_bullet(doc, "MCQ correction types are only allowed for single-letter MCQ answer keys.")
add_bullet(doc, "Large grammar-style answer keys are expanded in a bounded, deterministic way to prevent combinatorial growth.")

heading(doc, "Known character confusions", 2)
add_text(doc, "For corrected training rows, the pipeline counts directed changes from Captured to Published—for example, 0 → o or . → :. Repeated changes receive a smoothed frequency. For a new row, known_confusion_score is the strongest matching historical pattern; known_confusion_fraction is the share of required edits that are recognised.")
add_callout(doc, "Important", "A known confusion is evidence, not proof. It contributes to the review probability alongside confidence, similarity and text structure.")

add_callout(doc, "Guardrails", "Published is never a prediction feature; the parser does not inspect Captured when generating variations; and the model does not interpret the original image. The reviewer remains the final authority.", PALE_GOLD)

heading(doc, "3. How scoring works", 1)
add_text(doc, "The unreviewed source is data/hitl/no_human.csv. Scoring repeats the same feature calculations, applies the saved model, and writes a row-level review queue.")
add_table(doc, ["Stage", "What happens"], [
    ("1  Parse", "Generate bounded accepted variations from ANSWER Key only."),
    ("2  Measure", "Calculate confidence, similarities, shape and known-confusion evidence."),
    ("3  Predict", "Combine the signals into review_probability."),
    ("4  Flag", "Compare probability with the saved review threshold."),
    ("5  Prioritise", "Assign low, medium or high risk and suggest a correction type."),
], [1800, 7560])

heading(doc, "Risk-labelling rules", 2)
add_table(doc, ["Risk", "Rule", "Human-review flag"], [
    ("Low", "Probability is below review_threshold", "False"),
    ("Medium", "Probability is at or above review_threshold but below max(0.8, threshold)", "True"),
    ("High", "Probability is at or above max(0.8, review_threshold)", "True"),
], [1500, 5860, 2000])
add_text(doc, "The threshold is selected on validation data to meet the requested correction recall—95% by default—and can be overridden at scoring time. If the threshold is above 0.8, there is no medium band.")

heading(doc, "What the output contains", 2)
add_table(doc, ["Output", "How to use it"], [
    ("review_probability", "Primary ranking score; highest values first."),
    ("requires_human_review", "Main queue inclusion flag."),
    ("risk_label", "High/medium/low priority band."),
    ("predicted_correction_type", "Likely issue; supports triage but is not a correction."),
    ("similarity evidence", "Edit, character n-gram and token similarity for inspection."),
    ("known-confusion evidence", "Shows whether the edits resemble historical OCR patterns."),
], [2800, 6560])

add_page_break(doc)
heading(doc, "4. Recommended human-review workflow", 1)
doc.add_paragraph().paragraph_format.space_after = Pt(0)
add_number(doc, "Filter to requires_human_review = True.")
add_number(doc, "Sort by review_probability from highest to lowest; review high-risk rows first.")
add_number(doc, "Show the reviewer Captured, ANSWER Key, OCR confidence, evidence fields and—most importantly—the original scan image.")
add_number(doc, "Record the verified value and action in new fields; do not overwrite the original Captured value.")
add_number(doc, "Audit a small random sample of low-risk rows and a larger sample just below the threshold.")
add_number(doc, "Feed completed decisions back into the reviewed dataset and retrain on a controlled schedule.")

heading(doc, "Suggested reviewer fields", 2)
add_table(doc, ["Field", "Example values"], [
    ("human_review_status", "pending, completed, unable_to_determine"),
    ("human_corrected_value", "Verified response"),
    ("human_action", "confirmed_capture, corrected_capture, confirmed_blank, unreadable"),
    ("human_error_class", "substitution, insertion, deletion, multiple_edits, other"),
    ("reviewer / reviewed_at", "Audit trail"),
], [3000, 6360])

heading(doc, "Controls and safeguards", 2)
add_bullet(doc, "Never apply predicted correction types as automatic answer changes.")
add_bullet(doc, "Keep UID, scan_id, original Captured, model version and threshold for traceability.")
add_bullet(doc, "Monitor correction yield by probability band, question type, language group and PoS.")
add_bullet(doc, "Retrain when the OCR engine, question formats or language mix changes materially.")

heading(doc, "Leadership metrics", 2)
add_table(doc, ["Metric", "Leadership question answered"], [
    ("Correction recall", "What share of genuine corrections did we flag?"),
    ("Review precision", "What share of flagged rows actually needed correction?"),
    ("Review rate", "What share of all rows require human effort?"),
    ("Low-risk audit miss rate", "How often are apparently safe rows wrong?"),
    ("Reviewer time per row", "What is the operational cost?"),
], [3000, 6360])

add_page_break(doc)
heading(doc, "5. Limitations and next decisions", 1)
add_table(doc, ["Current limitation", "Recommended response"], [
    ("The model learns from historical reviewer behaviour.", "Check label consistency and document review policy."),
    ("Only 2.3% of rows are positive examples.", "Track recall and review yield; keep class weighting."),
    ("Answer-key expansion is bounded.", "Monitor complex keys and increase the cap only when runtime allows."),
    ("Known confusions are currently global.", "Consider separate tables for numeric, short-text and sentence responses once sample sizes are sufficient."),
    ("The model does not inspect images.", "Keep image review as the final authority."),
    ("Performance can drift.", "Review monthly/quarterly monitoring and establish retraining triggers."),
], [3900, 5460])

heading(doc, "Recommended rollout", 2)
add_callout(doc, "Phase 1 — Assisted review", "Send all high- and medium-risk rows to reviewers and audit low-risk samples. Measure recall, review precision, review rate and reviewer time.")
add_callout(doc, "Phase 2 — Controlled optimisation", "Adjust the review threshold based on review capacity and observed miss rate. Keep a fixed quality-control sample.")
add_callout(doc, "Phase 3 — Ongoing learning", "Add verified decisions to training data, version models and thresholds, and monitor drift by question and language segment.")

heading(doc, "Decision requested from leadership", 2)
add_bullet(doc, "Agree the acceptable missed-correction rate and target recall.")
add_bullet(doc, "Confirm available reviewer capacity and service-level expectations.")
add_bullet(doc, "Approve a low-risk sampling policy and model monitoring cadence.")
add_bullet(doc, "Define who owns model approval, threshold changes and retraining sign-off.")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT.resolve())
